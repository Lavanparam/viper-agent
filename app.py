from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel
from typing import List, Optional
from agent import GameTheoryAgent
import uvicorn
import os
from dotenv import load_dotenv
import io
from pypdf import PdfReader
from docx import Document
from fpdf import FPDF

load_dotenv()

app = FastAPI(title="Game Theory Agent")

# Mount static files for the frontend
app.mount("/static", StaticFiles(directory="static", html=True), name="static")

# Initialize Agent
try:
    agent = GameTheoryAgent()
except Exception as e:
    print(f"Failed to initialize agent: {e}")
    agent = None

class ChatMessage(BaseModel):
    role: str
    content: str

class ChatRequest(BaseModel):
    message: str
    history: List[ChatMessage]

class ExportRequest(BaseModel):
    history: List[ChatMessage]
    format: str

@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    content = ""
    try:
        if file.filename.endswith(".pdf"):
            reader = PdfReader(file.file)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    content += text + "\n"
        else:
            # Assume text-based
            content_bytes = await file.read()
            content = content_bytes.decode("utf-8", errors="ignore")
        
        return {"filename": file.filename, "content": content}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to process file: {str(e)}")

@app.post("/export")
async def export_chat(request: ExportRequest):
    try:
        if request.format == "pdf":
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            
            for msg in request.history:
                role = "User" if msg.role == "user" else "Viper"
                pdf.set_font("Arial", 'B', 12)
                pdf.cell(0, 10, f"{role}:", ln=True)
                pdf.set_font("Arial", size=12)
                # Handle unicode for FPDF (basic latin-1)
                safe_content = msg.content.encode('latin-1', 'replace').decode('latin-1') 
                pdf.multi_cell(0, 10, safe_content)
                pdf.ln(5)
            
            output = io.BytesIO()
            # fpdf1 output to string, encode to bytes
            pdf_str = pdf.output(dest='S')
            output.write(pdf_str.encode('latin-1'))
            output.seek(0)
            return StreamingResponse(output, media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=chat_history.pdf"})

        elif request.format == "docx":
            document = Document()
            document.add_heading('Viper Chat History', 0)
            
            for msg in request.history:
                role = "User" if msg.role == "user" else "Viper"
                document.add_heading(role, level=2)
                document.add_paragraph(msg.content)
            
            output = io.BytesIO()
            document.save(output)
            output.seek(0)
            pass # seek(0) done
            return StreamingResponse(output, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": "attachment; filename=chat_history.docx"})
        else:
            raise HTTPException(status_code=400, detail="Invalid format")

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")

@app.post("/chat")
async def chat_endpoint(request: ChatRequest):
    if not agent:
        raise HTTPException(status_code=500, detail="Agent not initialized (check API Key)")
    
    try:
        response_text = agent.generate_response(request.message, request.history)
        return {"response": response_text}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/")
async def read_root():
    return FileResponse('static/index.html')

if __name__ == "__main__":
    uvicorn.run(app, host="127.0.0.1", port=8000)
